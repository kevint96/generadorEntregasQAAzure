import streamlit as st
import os
import shutil
from zipfile import ZipFile
import zipfile
import tempfile
import subprocess
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
import inspect
import os
import xml.etree.ElementTree as ET
import gspread
import time  # Importar el módulo time
import logging
import re
import inspect
import ast
from datetime import datetime
import difflib
import glob
import base64
import sys
import xml.etree.ElementTree as ET
from collections import defaultdict
from lxml import etree
import json
import zlib
import urllib.parse
import requests
import concurrent.futures
import asyncio
import io
from datetime import date
import pandas as pd
import copy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX

# Ruta relativa a las plantillas
RUTA_BASE = os.path.join("plantillas", "plantilla_base.docx")
RUTA_MANUAL = os.path.join("plantillas", "plantilla_manual.docx")
RUTA_AUTORES = os.path.join("plantillas", "autores.txt")

# Al principio del script, asegúrate de inicializar una bandera
if "recargar_autores" not in st.session_state:
    st.session_state.recargar_autores = False

def print_with_line_number(msg):
    caller_frame = inspect.currentframe().f_back
    line_number = caller_frame.f_lineno
    st.success(f"Linea {line_number}: {msg}")
    print("")

def apply_format(run,fuente,size,negrita,color,highlight=None):
    run.font.name = fuente  # Cambiar el nombre de la fuente
    run.font.size = Pt(size)  # Cambiar el tamaño de la fuente
    run.font.bold = negrita  # Aplicar negrita
    run.font.color.rgb = RGBColor(0, 0, color)  # Cambiar el color del texto a rojo
    
    if highlight:  # Color de resaltado
        run.font.highlight_color = highlight

def replace_text_in_paragraph(paragraph, replacements):
    full_text = paragraph.text
    contador = 1
    ##st.success(f"Texto en linea: {full_text}")
    for key, value in replacements.items():
        if key in full_text:
            # Si el valor está vacío y la variable está sola, eliminar el párrafo completo
            if not value.strip() and full_text.strip() == key:
                p_element = paragraph._element
                parent = p_element.getparent()
                if parent is not None:
                    parent.remove(p_element)
                return

            # Si el valor está vacío pero la variable está entre más texto, simplemente eliminar la variable
            elif not value.strip():
                full_text = full_text.replace(key, "")
            else:
                full_text = full_text.replace(key, str(value))  # Actualiza full_text
                
                if key in '{acta}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Poppins Light',8,False,0)    # Aplicar formato al texto del párrafo
                    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                if key in '{nombre_servicio}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Poppins Light',8,False,0)  # Aplicar formato al texto del párrafo
                    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    
                if key in '{fecha_hoy}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Poppins Light',8,False,0)  # Aplicar formato al texto del párrafo
                    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    
                if key in '{nombre_autor}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Poppins Light',8,False,0)    # Aplicar formato al texto del párrafo
                    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                if key in '{num_hrv}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Poppins Light',8,False,0)    # Aplicar formato al texto del párrafo
                    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                if key in '{num_iniciativa}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Poppins Light',8,False,0)    # Aplicar formato al texto del párrafo
                    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    
                if key in '{NUM_INICIATIVA}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Poppins Light',8,False,0)    # Aplicar formato al texto del párrafo
                    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                if key in '{fecha_actual}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',8,False,0)    # Aplicar formato al texto del párrafo
                    #paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                if key in '{descripcion_ajuste}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',8,False,0)  # Aplicar formato al texto del párrafo
                
                if key in '{descripcion_pruebas_sugeridas}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                
                if key in '{proyecto_osb}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                
                if key in '{num_rel}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                    
                if key in '{contexto_ohs}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0,WD_COLOR_INDEX.YELLOW)  # Aplicar formato al texto del párrafo
                    
                if key in '{cksum}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                
                if key in '{fecha_azure}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                
                if key in '{branch}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,True,0)  # Aplicar formato al texto del párrafo
                
                if key in '{branch_git}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,True,0)  # Aplicar formato al texto del párrafo
                
                if key in '{operacion}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                
                if key in '{commit}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                
                if key in '{num_hrv2}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial MT',8,False,0)  # Aplicar formato al texto del párrafo
                
                if key in '{inicial_acta}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial MT',8,False,0)  # Aplicar formato al texto del párrafo
                    
                if key in '{nombre_servicio2}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial MT',8,False,0)  # Aplicar formato al texto del párrafo

                if key in '{num_iniciativa2}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial MT',8,False,0)  # Aplicar formato al texto del párrafo
                    
                if key in '{nombre_servicio_manual}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',14,True,0)  # Aplicar formato al texto del párrafo    
                
                if key in '{fecha_actual_manual}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo    
                    
                if key in '{endpoint}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',10,False,0,WD_COLOR_INDEX.YELLOW)  # Aplicar formato al texto del párrafo    

                if key in '{nombre_autor_manual}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo    
                
                if key in '{proyecto_osb_manual}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',10,False,0)  # Aplicar formato al texto del párrafo    

                if key in '{num_hrv_manual}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo
                    
                if key in '{nombre_servicio3}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo    
                
                if key in '{num_iniciativa_manual}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial',10,False,0)  # Aplicar formato al texto del párrafo 

                if key in '{bus}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                    
                if key in '{prueba}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                    
                if key in '{aut_puntual}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo

                if key in '{aut_prod}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Arial Narrow',8,False,0)  # Aplicar formato al texto del párrafo
                    
                if key in '{num_servicenow}':
                    paragraph.clear()  # Limpiar el párrafo
                    paragraph.add_run(full_text)  # Agregar el texto actualizado al párrafo
                    apply_format(paragraph.runs[0],'Poppins Light',8,False,0)  # Aplicar formato al texto del párrafo
            
def print_element_content(element, element_name):
    #st.success(f"Contenido del {element_name}:")
    for paragraph in element.paragraphs:
        print(paragraph.text)
    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)

def replace_text_in_element(element, replacements):
    for paragraph in element.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

def replace_text_in_doc(doc, replacements):
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        #st.success(f"Encabezado de la sección: {section.header}")
        #print_element_content(section.header, "Encabezado de la sección")
        replace_text_in_element(section.header, replacements)
        #st.success(f"Pie de página de la sección: {section.footer}")
        #print_element_content(section.footer, "Pie de página de la sección")
        replace_text_in_element(section.footer, replacements)
        # Agregamos este bloque específico para procesar las tablas dentro del encabezado de la sección 2
        if "Encabezado-Sección 2-" in [paragraph.text for paragraph in section.header.paragraphs]:
            # for table in section.header.tables:
                # for row in table.rows:
                    # for cell in row.cells:
                        # for paragraph in cell.paragraphs:
                            # print(paragraph.text)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, replacements)
    
    
    return doc


def reemplazar_variables(doc: Document, reemplazos: dict):
    def reemplazar_en_parrafo(parrafo):
        texto_original = ''.join(run.text for run in parrafo.runs)
        texto_nuevo = texto_original
        for key, value in reemplazos.items():
            texto_nuevo = texto_nuevo.replace(key, value)

        if texto_nuevo != texto_original:
            # Limpiar los runs existentes
            for i in range(len(parrafo.runs) - 1, -1, -1):
                parrafo._element.remove(parrafo.runs[i]._element)

            # Crear nuevo run con estilo similar
            nuevo_run = parrafo.add_run(texto_nuevo)
            nuevo_run.bold = False  # Puedes ajustarlo si quieres conservar estilos
            nuevo_run.italic = False

    for p in doc.paragraphs:
        reemplazar_en_parrafo(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    reemplazar_en_parrafo(p)

    return doc

def reemplazar_tabla_proyectos(doc: Document, proyectos_osb_filas, reemplazos_generales):
    from copy import deepcopy
    
    def estilo_disponible(doc, estilo_nombre):
        try:
            _ = doc.styles[estilo_nombre]
            return True
        except KeyError:
            return False

    estilo_lista = "ListBullet" if estilo_disponible(doc, "ListBullet") else "Normal"
    
   #print_with_line_number(f"reemplazos_generales: {reemplazos_generales}")
    total_tablas = len(doc.tables)
   #print_with_line_number(f"total_tablas: {total_tablas}")
    for table in doc.tables:
        #print_with_line_number(f"table: {table}")
        for i, row in enumerate(table.rows):
            #print_with_line_number(f"index: {i}")
            #print_with_line_number(f"row: {row}")
            for cell in row.cells:
                #print_with_line_number(f"cell.text: {cell.text}")
                if "{proyecto_osb}" in cell.text:
                   #print_with_line_number(f"cell.text: {cell.text}")
                    plantilla_row = table.rows[i]
                   #print_with_line_number(f"plantilla_row: {plantilla_row}")
                    table._tbl.remove(plantilla_row._tr)
                    for item in proyectos_osb_filas:
                        reemplazos = reemplazos_generales.copy()
                        reemplazos["{proyecto_osb}"] = item["proyecto_osb"]
                        reemplazos["{num_rel}"] = item["num_rel"]
                        reemplazos["{cksum}"] = item["cksum"]
                        reemplazos["{commit}"] = item["commit"]
                        reemplazos["{fecha_azure}"] = item["fecha_azure"]
                        # Copia profunda de la fila plantilla
                        new_tr = deepcopy(plantilla_row._tr)
                        table._tbl.append(new_tr)
                        new_row = table.rows[-1]

                        for j, cell_copy in enumerate(new_row.cells):
                            texto_base = plantilla_row.cells[j].text
                            for key, value in reemplazos.items():
                                texto_base = texto_base.replace(key, value)
                            
                            cell_copy.text = ""  # Limpiar contenido
                            p = cell_copy.paragraphs[0]
                            run = p.add_run(texto_base)
                            apply_format(run, fuente="Arial Narrow", size=10, negrita=False, color=0)
                    
                    break  # Solo salir de la fila actual

    # Reemplazo en párrafos: {proyecto_osb_lista}
    for i, paragraph in enumerate(doc.paragraphs):
        if "{proyecto_osb_lista}" in paragraph.text:
            p_index = i
            paragraph._element.getparent().remove(paragraph._element)

            for item in reversed(proyectos_osb_filas):
                new_p = doc.paragraphs[p_index].insert_paragraph_before(
                    f"{item['proyecto_osb']}.sbar", style=estilo_lista
                )
                run = new_p.runs[0]
                apply_format(run, fuente="Arial Narrow", size=8, negrita=False, color=0)
            break
    
    
def generar_documento(doc, nombre_resultado, reemplazos, proyectos_osb_filas=None):
    if proyectos_osb_filas:
        reemplazar_tabla_proyectos(doc, proyectos_osb_filas, reemplazos)
    
    doc_nuevo = replace_text_in_doc(doc, reemplazos)
    output_path = os.path.join(tempfile.gettempdir(), nombre_resultado)
    doc_nuevo.save(output_path)
    return output_path

def cargar_autores():
    if os.path.exists(RUTA_AUTORES):
        with open(RUTA_AUTORES, "r", encoding="utf-8") as f:
            autores = [line.strip() for line in f.readlines() if line.strip()]
    else:
        autores = []

    return autores

def guardar_autor(nuevo_autor):
    with open(RUTA_AUTORES, "a", encoding="utf-8") as f:
        f.write(f"{nuevo_autor.strip()}\n")


def main():
    st.set_page_config(layout="wide")
    
    if "num_hrv" not in st.session_state or st.session_state["num_hrv"].strip() == "":
        st.session_state["num_hrv"] = "XXXX"
    # Centrar título con HTML + CSS
    st.markdown(
        """
        <h1 style='text-align: center;'>
            Generador Entregas QA Azure 
        </h1>
        """,
        unsafe_allow_html=True
    )

    # Ajustar estilo visual
    st.markdown(
        """
        <style>
            /* Quitar los márgenes laterales del contenedor principal para que ocupe casi toda la pantalla */
            .reportview-container .main .block-container {
                max-width: 95vw;
                padding-left: 2rem;
                padding-right: 2rem;
            }
            
            .main .block-container {
                max-width: 95%;
                padding-left: 2rem;
                padding-right: 2rem;
            }
            input, textarea, select {
                font-size: 14px !important;
            }
        </style>
        """,
        unsafe_allow_html=True
    )
    

    # Fila 1
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        nombre_servicio = st.text_input("🛠️ Nombre del servicio")
    with col2:
        operacion = st.text_input("📡 Operación")
    with col3:
        # Estado inicial
        if "autores" not in st.session_state:
            st.session_state.autores = cargar_autores()

        # Agregar opción adicional
        opciones_autores = st.session_state.autores + ["📝 Agregar nuevo..."]

        # Combo de selección
        nombre_autor = st.selectbox("👤 Nombre del autor", cargar_autores())

        # # Si se escoge agregar uno nuevo
        # if nombre_autor == "📝 Agregar nuevo...":
            # nuevo_autor = st.text_input("✍️ Escribe el nuevo autor y presiona Enter:")

            # if nuevo_autor.strip() != "":
                # if nuevo_autor.strip() not in st.session_state.autores:
                    # guardar_autor(nuevo_autor.strip())
                    # st.session_state.autores.append(nuevo_autor.strip())
                    # st.success(f"✅ Autor '{nuevo_autor.strip()}' agregado correctamente.")
                # else:
                    # st.warning("⚠️ El autor ya existe.")
    with col4:
        bus = st.selectbox("💻 BUS", ["Otorgamiento", "Digital"])

    # Fila 2
    col5, col6, col7, col8 = st.columns(4)
    with col5:
        num_iniciativa = st.text_input("🆔 Número Iniciativa") 
    with col6:
        num_servicenow = st.text_input("🆔 Número Servicenow")
    with col7:
        num_hrv = st.text_input("🔢 Número Harvest", value=st.session_state["num_hrv"])
    with col8:
        st.text_input("🧬 Consecutivo", value=1, disabled=True)
        #commit = st.text_input("🧬 Commit")
        
    # Fila 3
    col9, col10, col11, col12 = st.columns(4)
    with col9:
        #st.write("✅ Autorizaciones ➡️")
        tipo_acta = st.checkbox("Catalogación Azure")
        #st.selectbox("💻 Tipo Acta", ["Harvest", "Azure"])
        #st.subheader("✅ Autorizaciones")
        # st.text_input("🛠️ Nombre del servicio", value=valor, disabled=True)
    with col10:
        prueba = st.checkbox("📡 Pruebas (Certificación)")
    with col11:
        aut_puntual = st.checkbox("📡 Con aut. (Validación puntual QA)")
    with col12:
        aut_prod = st.checkbox("📡 Con aut. (Producción)")
    
    # Construir la variable Acta
    inicial_acta = "MW"
    bo = "_BO" if bus == "Otorgamiento" else ""
    if tipo_acta:
        inicial_acta ="AMW"
    id_iniciativa = num_iniciativa if num_iniciativa.strip() else num_servicenow
    
    if num_hrv.strip():
        num_hrv = num_hrv
    else:
        if tipo_acta:
            num_hrv = "XXXXX"
        else:
            num_hrv = "XXXX"
    acta = f"{inicial_acta}{num_hrv}_OSB12C{bo}_{nombre_servicio}_ID_{id_iniciativa}_1"
    
    # # Mostrar Acta no editable con HTML (readonly)
    # st.markdown("### 📝 Acta generada")
    # st.markdown(f"""
        # <input type="text" value="{acta}" readonly style="width: 100%; padding: 0.5em; font-size: 1em;" />
    # """, unsafe_allow_html=True)
    st.text_input("📝 Acta", value=acta, disabled=True)
    # Fila exclusiva para branch
    
    # Lógica para establecer el valor por defecto
    if num_servicenow.strip():
        default_index = 1  # hotfix
    elif num_iniciativa.strip():
        default_index = 0  # feature
    else:
        default_index = 0  # default por si ambos están vacíos
    
    branch = st.selectbox("🌱 Branch", ["feature", "hotfix"], index=default_index, disabled=True)
    
    branch_completo = f"{branch}/{inicial_acta}{num_hrv}_OSB12C{bo}_{nombre_servicio}_ID_{id_iniciativa}"
    
    branch_git = st.text_input("🌴 Branch git", value=branch_completo, disabled=True)
    
    checkout = f"git checkout -b {branch_completo} origin/{branch_completo}"
    
    st.text_input("🔀 Checkout branch", value=checkout, disabled=True)
    
    # Tabla editable de proyectos
    st.markdown("### 🧩 Proyectos OSB (máximo 4 - En orden de instalación)")
    import pandas as pd

    proyectos_default = pd.DataFrame({
        "Proyecto OSB": ["", "", "", ""],
        "Release": ["", "", "", ""],
        "Checksum": ["", "", "", ""],
        "Commit": ["", "", "", ""],
        "Fecha Azure": ["", "", "", ""]
    })

    proyectos_input = st.data_editor(
        proyectos_default,
        num_rows="dynamic",
        use_container_width=True,
        key="proyectos_osb_input"
    )

    proyectos_osb = [
        {
            "proyecto_osb": row["Proyecto OSB"].strip(),
            "num_rel": str(row["Release"]).strip(),
            "cksum": str(row["Checksum"]).strip(),
            "commit": str(row["Commit"]).strip(),
            "fecha_azure": str(row["Fecha Azure"]).strip()
        }
        for _, row in proyectos_input.iterrows()
        if row["Proyecto OSB"] and row["Release"] and row["Checksum"] and row["Commit"] and row["Fecha Azure"]
    ][:4]

    # Descripción + archivos
    descripcion_ajuste = st.text_area("📝 Descripción funcional del ajuste")
    
    descripcion_pruebas_sugeridas = st.text_area("📝 Descripción pruebas sugeridas")
    
    nuevo_endpoint = st.checkbox("Nuevo endpoint OHS")
    
    if not nuevo_endpoint:
        endpoint = 'N/A'
        contexto_ohs = ""
    else:
        endpoint = st.text_input("🛠️ Url OHS")
        contexto_ohs = f"Agregar el nuevo contexto:\n {endpoint} en el ambiente de {bus}"

    # Carga directa (sin subir)
    plantilla_doc = Document(RUTA_BASE)
    plantilla_manual = Document(RUTA_MANUAL)
    # col_file1, col_file2 = st.columns(2)
    # with col_file1:
        # plantilla_doc = st.file_uploader("📎 Plantilla base (documento principal)", type="docx")
    # with col_file2:
        # plantilla_manual = st.file_uploader("📎 Plantilla manual instalación", type="docx")

    #submit = st.form_submit_button("📄 Generar documentos")

    if st.button("📄 Generar documentos"):
        if not plantilla_doc:
            st.error("❌ Por favor, suba la plantilla base (documento principal).")
        elif not plantilla_manual:
            st.error("❌ Por favor, suba la plantilla manual de instalación.")
        elif not proyectos_osb:
            st.error("❌ Por favor, ingrese al menos un proyecto OSB válido.")
        elif not nombre_servicio.strip() or not operacion.strip():
            st.error("❌ Por favor ingrese el nombre del servicio y operación")
        elif not nombre_autor.strip() or not id_iniciativa.strip() or not num_hrv.strip():
            st.error("❌ Por favor ingrese el nombre del autor, iniciativa y numero harvest")
        elif not(prueba or aut_puntual or aut_prod):
            st.error("❌ Por favor seleccione el tipo de autorización")
        else:
            fecha_actual = date.today().strftime("%Y-%m-%d")
            fecha_hoy = date.today().strftime("%d/%m/%Y")
            fecha_azure = date.today().strftime("%Y%m%d")
            
            proyecto_osb = proyectos_osb[0]["proyecto_osb"] if proyectos_osb else ""
            num_rel = proyectos_osb[0]["num_rel"] if proyectos_osb else ""
            cksum = proyectos_osb[0]["cksum"] if proyectos_osb else ""
            commit = proyectos_osb[0]["commit"] if proyectos_osb else ""
            fecha_azure = proyectos_osb[0]["fecha_azure"] if proyectos_osb else date.today().strftime("%Y%m%d")+'.1'
            prueba = "X" if prueba else ""
            aut_puntual = "X" if aut_puntual else ""
            aut_prod = "X" if aut_prod else ""
            
            #print_with_line_number(f"prueba: {prueba}")
            #print_with_line_number(f"aut_puntual: {aut_puntual}")
            #print_with_line_number(f"aut_prod: {aut_prod}")

            reemplazos = {
                "{fecha_actual}": fecha_actual,
                "{fecha_hoy}": fecha_hoy,
                "{fecha_azure}": fecha_azure,
                "{nombre_servicio}": nombre_servicio,
                "{nombre_servicio_manual}": nombre_servicio,
                "{fecha_actual_manual}": fecha_actual,
                "{nombre_autor_manual}": nombre_autor,
                "{proyecto_osb_manual}": proyecto_osb,
                "{num_hrv_manual}": num_hrv,
                "{nombre_servicio3}": nombre_servicio,
                "{num_iniciativa_manual}": id_iniciativa,
                "{nombre_autor}": nombre_autor,
                "{num_hrv}": num_hrv,
                "{NUM_INICIATIVA}": id_iniciativa,
                "{num_iniciativa}": num_iniciativa,
                "{num_servicenow}": num_servicenow,
                "{bus}": bus,
                "{prueba}": prueba,
                "{aut_puntual}": aut_puntual,
                "{aut_prod}": aut_prod,
                "{num_hrv2}": num_hrv,
                "{inicial_acta}": inicial_acta,
                "{nombre_servicio2}": nombre_servicio,
                "{num_iniciativa2}": id_iniciativa,
                "{descripcion_ajuste}": descripcion_ajuste,
                "{descripcion_pruebas_sugeridas}": descripcion_pruebas_sugeridas,
                "{proyecto_osb}": proyecto_osb,
                "{proyecto_osb_lista}": proyecto_osb,
                "{operacion}": operacion,
                "{commit}": commit,
                "{num_rel}": num_rel,
                "{cksum}": cksum,
                "{branch}": branch,
                "{endpoint}": endpoint,
                "{branch_git}": branch_git,
                "{contexto_ohs}": contexto_ohs,
                "{acta}": acta
            }

            nombre_doc = f"{acta}.docx"
            nombre_manual = f"Manual_Instalacion_OSB12C{bo}_{nombre_servicio}.docx"

            path_out_doc = generar_documento(plantilla_doc, nombre_doc, reemplazos, proyectos_osb)
            path_out_manual = generar_documento(plantilla_manual, nombre_manual, reemplazos, proyectos_osb)

            # Nombre de carpeta interna dentro del .zip
            carpeta_zip = f"{inicial_acta}{num_hrv}-{nombre_servicio}"

            # Crear zip en memoria
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                zipf.write(path_out_doc, arcname=f"{carpeta_zip}/{nombre_doc}")
                zipf.write(path_out_manual, arcname=f"{carpeta_zip}/{nombre_manual}")

            st.success(f"📄 Documentos generados: ✅ {carpeta_zip}")
            # Mostrar botón para descargar el zip
            st.download_button(
                label="📦 Descargar documentos (ZIP)",
                data=zip_buffer.getvalue(),
                file_name=f"{carpeta_zip}.zip",
                mime="application/zip"
            )

if __name__ == "__main__":
    main()
